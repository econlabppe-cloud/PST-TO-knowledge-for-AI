from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def write_text_docx(path: Path, text: str, *, title: str | None = None) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    if title:
        document.add_heading(title, level=1)

    for line in text.splitlines():
        if not line.strip():
            document.add_paragraph("")
            continue

        paragraph = document.add_paragraph()
        if _is_heading_line(line):
            run = paragraph.add_run(line)
            run.bold = True
        elif line.startswith("- "):
            paragraph.style = "List Bullet"
            paragraph.add_run(line[2:])
        else:
            paragraph.add_run(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _is_heading_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith(("TOPIC:", "PERSON:", "EMAIL", "NOTEBOOKLM", "CLASSIFICATION RULES", "TOPICS", "PEOPLE")):
        return True
    if stripped.startswith(("Date range:", "Emails:", "Email count:", "Total ", "Representative emails:")):
        return True
    if stripped.startswith(("PART:", "TOP SENDERS:", "TOP SUBTOPICS:", "TOP PEOPLE:", "TOP TOPICS:")):
        return True
    return False
